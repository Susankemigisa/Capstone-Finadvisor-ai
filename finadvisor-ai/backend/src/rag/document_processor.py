"""
Document processor — ingests uploaded files into the RAG pipeline.

Supports: PDF, DOCX, TXT, MD, CSV
Pipeline per file:
    1. Extract raw text from the file bytes
    2. Split into overlapping chunks (size + overlap from settings)
    3. Return chunk dicts ready for embedding + storage

Chunk dict shape:
    {
        "content":     str,   # the chunk text
        "chunk_index": int,   # position in the document (0-based)
        "char_start":  int,   # character offset in original text
        "char_end":    int,
        "metadata":    dict,  # filename, file_type, page (PDF only)
    }
"""

from __future__ import annotations

import io
import csv
from typing import BinaryIO

from langchain_text_splitters import RecursiveCharacterTextSplitter

from src.config.settings import settings
from src.utils.logger import get_logger

logger = get_logger(__name__)

# Supported MIME types and their processors
SUPPORTED_TYPES = {
    "application/pdf":                                         "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain":                                              "txt",
    "text/markdown":                                           "txt",
    "text/csv":                                                "csv",
}

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".csv"}

# Absolute limits — protect embedding cost and context window
MAX_FILE_SIZE_MB  = 20
MAX_CHUNKS        = 500
MIN_CHUNK_CHARS   = 50   # discard tiny fragments


def process_document(
    file_bytes:  bytes,
    filename:    str,
    content_type: str = "",
) -> list[dict]:
    """
    Main entry point — process a file and return its chunks.

    Args:
        file_bytes:    Raw file content
        filename:      Original filename (used for metadata + type detection)
        content_type:  MIME type (optional, falls back to extension detection)

    Returns:
        List of chunk dicts ready for embedding.
        Returns [] on unsupported type or extraction failure.

    Raises:
        ValueError: if the file exceeds MAX_FILE_SIZE_MB
    """
    size_mb = len(file_bytes) / (1024 * 1024)
    if size_mb > MAX_FILE_SIZE_MB:
        raise ValueError(
            f"File '{filename}' is {size_mb:.1f} MB — maximum allowed is {MAX_FILE_SIZE_MB} MB."
        )

    file_type = _detect_type(filename, content_type)
    if not file_type:
        logger.warning("unsupported_file_type", filename=filename, content_type=content_type)
        return []

    logger.info("processing_document", filename=filename, file_type=file_type, size_mb=round(size_mb, 2))

    # Extract raw text
    try:
        raw_text = _extract_text(file_bytes, file_type, filename)
    except Exception as e:
        logger.error("text_extraction_failed", filename=filename, error=str(e))
        return []

    if not raw_text or len(raw_text.strip()) < MIN_CHUNK_CHARS:
        logger.warning("document_empty_after_extraction", filename=filename)
        return []

    # Split into chunks
    chunks = _split_text(
        text=raw_text,
        filename=filename,
        file_type=file_type,
    )

    logger.info("document_processed", filename=filename, chunks=len(chunks))
    return chunks[:MAX_CHUNKS]


def _detect_type(filename: str, content_type: str) -> str | None:
    """Detect file type from MIME type or extension. Returns short type string."""
    # Try MIME type first
    if content_type and content_type in SUPPORTED_TYPES:
        return SUPPORTED_TYPES[content_type]

    # Fall back to extension
    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    ext_map = {".pdf": "pdf", ".docx": "docx", ".txt": "txt", ".md": "txt", ".csv": "csv"}
    return ext_map.get(ext)


def _extract_text(file_bytes: bytes, file_type: str, filename: str) -> str:
    """Dispatch to the correct extractor for the file type."""
    if file_type == "pdf":
        return _extract_pdf(file_bytes)
    if file_type == "docx":
        return _extract_docx(file_bytes)
    if file_type == "txt":
        return file_bytes.decode("utf-8", errors="replace")
    if file_type == "csv":
        return _extract_csv(file_bytes)
    return ""


def _extract_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pypdf."""
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"[Page {i + 1}]\n{text}")
    return "\n\n".join(pages)


def _extract_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract table cells
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return "\n\n".join(paragraphs)


def _extract_csv(file_bytes: bytes) -> str:
    """Convert CSV rows to readable text blocks."""
    text = file_bytes.decode("utf-8", errors="replace")
    reader = csv.DictReader(io.StringIO(text))
    lines = []
    try:
        for i, row in enumerate(reader):
            if i >= 1000:   # Cap at 1000 rows
                lines.append("[... truncated at 1000 rows ...]")
                break
            row_text = ", ".join(f"{k}: {v}" for k, v in row.items() if v and v.strip())
            if row_text:
                lines.append(row_text)
    except Exception:
        # Fallback: treat as plain text
        return text
    return "\n".join(lines)


def _split_text(text: str, filename: str, file_type: str) -> list[dict]:
    """
    Split raw text into overlapping chunks using LangChain's
    RecursiveCharacterTextSplitter.

    Chunk size and overlap come from settings so they can be tuned
    without code changes.
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.CHUNK_SIZE,
        chunk_overlap=settings.CHUNK_OVERLAP,
        length_function=len,
        separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""],
    )

    raw_chunks = splitter.split_text(text)

    chunks = []
    char_offset = 0
    for i, chunk_text in enumerate(raw_chunks):
        chunk_text = chunk_text.strip()
        if len(chunk_text) < MIN_CHUNK_CHARS:
            continue

        char_start = text.find(chunk_text, char_offset)
        char_end   = char_start + len(chunk_text) if char_start >= 0 else -1
        if char_start >= 0:
            char_offset = char_start + 1

        chunks.append({
            "content":     chunk_text,
            "chunk_index": i,
            "char_start":  max(char_start, 0),
            "char_end":    max(char_end,   0),
            "metadata": {
                "filename":  filename,
                "file_type": file_type,
            },
        })

    return chunks
